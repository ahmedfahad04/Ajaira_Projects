from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class WordDocumentManager:
    def __init__(self, document_path):
        self.document_path = document_path

    def extract_text(self):
        document = Document(self.document_path)
        content_lines = [paragraph.text for paragraph in document.paragraphs]
        return '\n'.join(content_lines)

    def insert_text(self, content, font_size=12, alignment='left'):
        try:
            document = Document()
            new_paragraph = document.add_paragraph()
            new_paragraph.add_run(content).font.size = Pt(font_size)
            alignment_value = self._determine_alignment(alignment)
            new_paragraph.alignment = alignment_value
            document.save(self.document_path)
            return True
        except:
            return False

    def incorporate_heading(self, heading, level=1):
        try:
            document = Document(self.document_path)
            document.add_heading(heading, level)
            document.save(self.document_path)
            return True
        except:
            return False

    def incorporate_table(self, table_data):
        try:
            document = Document(self.document_path)
            table = document.add_table(rows=len(table_data), cols=len(table_data[0]))
            for row_idx, row in enumerate(table_data):
                for col_idx, value in enumerate(row):
                    table.cell(row_idx, col_idx).text = str(value)
            document.save(self.document_path)
            return True
        except:
            return False

    def _determine_alignment(self, alignment):
        alignment_options = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
        }
        return alignment_options.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)
