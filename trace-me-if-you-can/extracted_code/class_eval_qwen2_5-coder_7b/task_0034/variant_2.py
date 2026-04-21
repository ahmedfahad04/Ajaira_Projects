from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class DocxHandler:
    def __init__(self, file_location):
        self.file_location = file_location

    def get_content(self):
        doc = Document(self.file_location)
        content = '\n'.join([para.text for para in doc.paragraphs])
        return content

    def write_content(self, text, size=12, align='left'):
        try:
            doc = Document()
            para = doc.add_paragraph()
            para.add_run(text).font.size = Pt(size)
            para.alignment = getattr(WD_PARAGRAPH_ALIGNMENT, align.capitalize(), WD_PARAGRAPH_ALIGNMENT.LEFT)
            doc.save(self.file_location)
            return True
        except:
            return False

    def insert_heading(self, heading, level=1):
        try:
            doc = Document(self.file_location)
            doc.add_heading(heading, level)
            doc.save(self.file_location)
            return True
        except:
            return False

    def add_matrix(self, matrix_data):
        try:
            doc = Document(self.file_location)
            table = doc.add_table(rows=len(matrix_data), cols=len(matrix_data[0]))
            for row_idx, row in enumerate(matrix_data):
                for col_idx, value in enumerate(row):
                    table.cell(row_idx, col_idx).text = str(value)
            doc.save(self.file_location)
            return True
        except:
            return False

    def _align_value(self, alignment):
        alignment_map = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
        }
        return alignment_map.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)
