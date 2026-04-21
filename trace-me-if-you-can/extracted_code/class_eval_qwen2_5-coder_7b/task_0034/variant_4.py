from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class DocxModifier:
    def __init__(self, file_path):
        self.file_path = file_path

    def read_content(self):
        doc = Document(self.file_path)
        content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        return content

    def write_content(self, text, size=12, align='left'):
        try:
            doc = Document()
            para = doc.add_paragraph()
            para.add_run(text).font.size = Pt(size)
            para.alignment = getattr(WD_PARAGRAPH_ALIGNMENT, align.capitalize(), WD_PARAGRAPH_ALIGNMENT.LEFT)
            doc.save(self.file_path)
            return True
        except:
            return False

    def insert_heading(self, heading, level=1):
        try:
            doc = Document(self.file_path)
            doc.add_heading(heading, level)
            doc.save(self.file_path)
            return True
        except:
            return False

    def add_matrix(self, matrix_data):
        try:
            doc = Document(self.file_path)
            table = doc.add_table(rows=len(matrix_data), cols=len(matrix_data[0]))
            for i, row in enumerate(matrix_data):
                for j, value in enumerate(row):
                    table.cell(i, j).text = str(value)
            doc.save(self.file_path)
            return True
        except:
            return False

    def _determine_alignment(self, alignment):
        alignment_options = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
        }
        return alignment_options.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)
