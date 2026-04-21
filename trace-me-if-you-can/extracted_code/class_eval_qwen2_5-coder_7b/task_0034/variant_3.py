from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class DocxProcessor:
    def __init__(self, doc_path):
        self.doc_path = doc_path

    def load_text(self):
        doc = Document(self.doc_path)
        return '\n'.join([para.text for para in doc.paragraphs])

    def save_text(self, content, font_size=12, alignment='left'):
        try:
            doc = Document()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(content)
            run.font.size = Pt(font_size)
            paragraph.alignment = getattr(WD_PARAGRAPH_ALIGNMENT, alignment.capitalize(), WD_PARAGRAPH_ALIGNMENT.LEFT)
            doc.save(self.doc_path)
            return True
        except:
            return False

    def add_section(self, heading, level=1):
        try:
            doc = Document(self.doc_path)
            doc.add_heading(heading, level)
            doc.save(self.doc_path)
            return True
        except:
            return False

    def include_table(self, table_data):
        try:
            doc = Document(self.doc_path)
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            for i, row in enumerate(table_data):
                for j, value in enumerate(row):
                    table.cell(i, j).text = str(value)
            doc.save(self.doc_path)
            return True
        except:
            return False

    def _get_alignment_code(self, alignment):
        alignment_options = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
        }
        return alignment_options.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)
