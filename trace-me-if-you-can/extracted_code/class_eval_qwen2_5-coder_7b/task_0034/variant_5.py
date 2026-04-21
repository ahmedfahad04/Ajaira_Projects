from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class DocxEditor:
    def __init__(self, file_location):
        self.file_location = file_location

    def extract_text(self):
        doc = Document(self.file_location)
        content = '\n'.join([para.text for para in doc.paragraphs])
        return content

    def insert_text(self, content, font_size=12, alignment='left'):
        try:
            doc = Document()
            para = doc.add_paragraph()
            para.add_run(content).font.size = Pt(font_size)
            alignment_value = self._determine_alignment(alignment)
            para.alignment = alignment_value
            doc.save(self.file_location)
            return True
        except:
            return False

    def include_heading(self, heading, level=1):
        try:
            doc = Document(self.file_location)
            doc.add_heading(heading, level)
            doc.save(self.file_location)
            return True
        except:
            return False

    def add_table(self, table_data):
        try:
            doc = Document(self.file_location)
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            for i, row in enumerate(table_data):
                for j, value in enumerate(row):
                    table.cell(i, j).text = str(value)
            doc.save(self.file_location)
            return True
        except:
            return False

    def _get_alignment_code(self, alignment):
        alignment_options = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
        }
        return alignment_options.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)
