from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocumentOperations:
    @staticmethod
    def extract_text(file_path):
        doc = Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])

    @staticmethod
    def create_document_with_text(file_path, content, font_size, alignment):
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(content)
        run.font.size = Pt(font_size)
        paragraph.alignment = DocumentOperations.get_alignment_enum(alignment)
        doc.save(file_path)

    @staticmethod
    def append_heading_to_document(file_path, heading, level):
        doc = Document(file_path)
        doc.add_heading(heading, level)
        doc.save(file_path)

    @staticmethod
    def append_table_to_document(file_path, data):
        doc = Document(file_path)
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        for i, row in enumerate(data):
            for j, cell_value in enumerate(row):
                table.cell(i, j).text = str(cell_value)
        doc.save(file_path)

    @staticmethod
    def get_alignment_enum(alignment):
        alignment_map = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
        }
        return alignment_map.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)


class DocFileHandler:
    def __init__(self, file_path):
        self.file_path = file_path
        self.operations = DocumentOperations()

    def read_text(self):
        return self.operations.extract_text(self.file_path)

    def write_text(self, content, font_size=12, alignment='left'):
        try:
            self.operations.create_document_with_text(self.file_path, content, font_size, alignment)
            return True
        except:
            return False

    def add_heading(self, heading, level=1):
        try:
            self.operations.append_heading_to_document(self.file_path, heading, level)
            return True
        except:
            return False

    def add_table(self, data):
        try:
            self.operations.append_table_to_document(self.file_path, data)
            return True
        except:
            return False

    def _get_alignment_value(self, alignment):
        return self.operations.get_alignment_enum(alignment)
