from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocFileHandler:
    def __init__(self, file_path):
        self.file_path = file_path
        self._operations = {
            'write': self._execute_write,
            'add_heading': self._execute_add_heading,
            'add_table': self._execute_add_table
        }

    def read_text(self):
        doc = Document(self.file_path)
        paragraphs = [p.text for p in doc.paragraphs]
        return "\n".join(paragraphs)

    def write_text(self, content, font_size=12, alignment='left'):
        return self._execute_operation('write', content=content, font_size=font_size, alignment=alignment)

    def add_heading(self, heading, level=1):
        return self._execute_operation('add_heading', heading=heading, level=level)

    def add_table(self, data):
        return self._execute_operation('add_table', data=data)

    def _execute_operation(self, operation_type, **kwargs):
        try:
            self._operations[operation_type](**kwargs)
            return True
        except:
            return False

    def _execute_write(self, content, font_size, alignment):
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(content)
        run.font.size = Pt(font_size)
        paragraph.alignment = self._get_alignment_value(alignment)
        doc.save(self.file_path)

    def _execute_add_heading(self, heading, level):
        doc = Document(self.file_path)
        doc.add_heading(heading, level)
        doc.save(self.file_path)

    def _execute_add_table(self, data):
        doc = Document(self.file_path)
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        list(map(lambda item: self._set_cell_text(table, item[0], item[1]), 
                 [(i, j, data[i][j]) for i in range(len(data)) for j in range(len(data[i]))]))
        doc.save(self.file_path)

    def _set_cell_text(self, table, row_col_value):
        i, j, value = row_col_value
        table.cell(i, j).text = str(value)

    def _get_alignment_value(self, alignment):
        alignments = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
        }
        return alignments.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)
