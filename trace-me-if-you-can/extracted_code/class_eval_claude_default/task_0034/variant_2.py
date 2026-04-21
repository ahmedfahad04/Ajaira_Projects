from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from functools import wraps


def safe_document_operation(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        try:
            func(*args, **kwargs)
            return True
        except:
            return False
    return wrapper


class DocFileHandler:
    def __init__(self, file_path):
        self.file_path = file_path
        self._alignment_mapping = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
        }

    def read_text(self):
        with Document(self.file_path) as doc:
            text_parts = [paragraph.text for paragraph in doc.paragraphs]
        return "\n".join(text_parts)

    @safe_document_operation
    def write_text(self, content, font_size=12, alignment='left'):
        doc = Document()
        paragraph = doc.add_paragraph()
        self._configure_paragraph_run(paragraph, content, font_size)
        paragraph.alignment = self._resolve_alignment(alignment)
        doc.save(self.file_path)

    @safe_document_operation
    def add_heading(self, heading, level=1):
        doc = Document(self.file_path)
        doc.add_heading(heading, level)
        doc.save(self.file_path)

    @safe_document_operation
    def add_table(self, data):
        doc = Document(self.file_path)
        table = self._create_populated_table(doc, data)
        doc.save(self.file_path)

    def _configure_paragraph_run(self, paragraph, content, font_size):
        run = paragraph.add_run(content)
        run.font.size = Pt(font_size)

    def _create_populated_table(self, doc, data):
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        for row_idx, row_data in enumerate(data):
            for col_idx, cell_data in enumerate(row_data):
                table.cell(row_idx, col_idx).text = str(cell_data)
        return table

    def _resolve_alignment(self, alignment):
        return self._alignment_mapping.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)

    def _get_alignment_value(self, alignment):
        return self._resolve_alignment(alignment)
