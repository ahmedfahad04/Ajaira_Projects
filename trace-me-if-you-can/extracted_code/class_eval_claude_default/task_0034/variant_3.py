from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from contextlib import contextmanager


class DocFileHandler:
    def __init__(self, file_path):
        self.file_path = file_path

    def read_text(self):
        doc = Document(self.file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return "\n".join(text)

    def write_text(self, content, font_size=12, alignment='left'):
        with self._error_handling():
            doc = Document()
            self._create_formatted_paragraph(doc, content, font_size, alignment)
            doc.save(self.file_path)
            return True
        return False

    def add_heading(self, heading, level=1):
        with self._error_handling():
            with self._document_context() as doc:
                doc.add_heading(heading, level)
                doc.save(self.file_path)
            return True
        return False

    def add_table(self, data):
        with self._error_handling():
            with self._document_context() as doc:
                rows, cols = len(data), len(data[0])
                table = doc.add_table(rows=rows, cols=cols)
                self._populate_table_cells(table, data)
                doc.save(self.file_path)
            return True
        return False

    @contextmanager
    def _error_handling(self):
        try:
            yield
        except:
            pass

    @contextmanager
    def _document_context(self):
        doc = Document(self.file_path)
        yield doc

    def _create_formatted_paragraph(self, doc, content, font_size, alignment):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(content)
        font = run.font
        font.size = Pt(font_size)
        alignment_value = self._get_alignment_value(alignment)
        paragraph.alignment = alignment_value

    def _populate_table_cells(self, table, data):
        for i, row in enumerate(data):
            for j, cell_value in enumerate(row):
                table.cell(i, j).text = str(cell_value)

    def _get_alignment_value(self, alignment):
        alignment_options = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
        }
        return alignment_options.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)
