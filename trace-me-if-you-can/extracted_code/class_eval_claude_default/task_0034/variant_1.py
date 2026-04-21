from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class DocFileHandler:
    ALIGNMENT_MAP = {
        'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
        'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
        'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
    }
    
    def __init__(self, file_path):
        self.file_path = file_path

    def read_text(self):
        doc = Document(self.file_path)
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)

    def write_text(self, content, font_size=12, alignment='left'):
        return self._safe_document_operation(self._write_text_operation, content, font_size, alignment)

    def add_heading(self, heading, level=1):
        return self._safe_document_operation(self._add_heading_operation, heading, level)

    def add_table(self, data):
        return self._safe_document_operation(self._add_table_operation, data)

    def _safe_document_operation(self, operation, *args):
        try:
            operation(*args)
            return True
        except:
            return False

    def _write_text_operation(self, content, font_size, alignment):
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(content)
        run.font.size = Pt(font_size)
        paragraph.alignment = self.ALIGNMENT_MAP.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)
        doc.save(self.file_path)

    def _add_heading_operation(self, heading, level):
        doc = Document(self.file_path)
        doc.add_heading(heading, level)
        doc.save(self.file_path)

    def _add_table_operation(self, data):
        doc = Document(self.file_path)
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        for i, row in enumerate(data):
            for j, cell_value in enumerate(row):
                table.cell(i, j).text = str(cell_value)
        doc.save(self.file_path)

    def _get_alignment_value(self, alignment):
        return self.ALIGNMENT_MAP.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)
